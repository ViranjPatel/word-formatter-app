from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls, qn
from docx.enum.dml import MSO_THEME_COLOR
import tempfile
import os
import re
from collections import defaultdict, Counter
from functools import lru_cache
import logging

# Configure logging for optional debug output
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)

class LaTeXConverter:
    """Advanced Word to LaTeX converter with comprehensive formatting support"""
    
    def __init__(self, debug=False):
        self.debug = debug
        self.latex_content = []
        self.packages = set()
        self.document_class = "article"
        self.extracted_styles = {}
        self.font_families_used = set()
        self.colors_used = set()
        self.custom_commands = []
        
        # Enhanced style mappings
        self.style_mappings = {
            'Title': r'\title{{{content}}}',
            'Subtitle': r'\subtitle{{{content}}}',
            'Heading 1': r'\section{{{content}}}',
            'Heading 2': r'\subsection{{{content}}}',
            'Heading 3': r'\subsubsection{{{content}}}',
            'Heading 4': r'\paragraph{{{content}}}',
            'Heading 5': r'\subparagraph{{{content}}}',
            'Heading 6': r'\subparagraph{{{content}}}',
        }
        
        # Alignment mappings
        self.alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: 'flushleft',
            WD_ALIGN_PARAGRAPH.CENTER: 'center',
            WD_ALIGN_PARAGRAPH.RIGHT: 'flushright',
            WD_ALIGN_PARAGRAPH.JUSTIFY: 'justify',
        }
        
        if debug:
            logger.setLevel(logging.INFO)
    
    def convert_document(self, docx_path):
        """Convert Word document to LaTeX with comprehensive formatting"""
        doc = Document(docx_path)
        
        # Extract comprehensive document information
        self._analyze_document_styles(doc)
        self._extract_document_properties(doc)
        
        # Initialize LaTeX document with proper packages
        self._initialize_comprehensive_latex()
        
        # Process document content with advanced formatting
        self._process_document_content(doc)
        
        # Finalize document
        self._finalize_latex_document()
        
        # Generate output file
        output_path = tempfile.mktemp(suffix='.tex')
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(self.latex_content))
        
        if self.debug:
            logger.info(f"Advanced LaTeX document created: {output_path}")
        
        return output_path
    
    def _analyze_document_styles(self, doc):
        """Comprehensively analyze all styles used in the document"""
        # Extract all paragraph and character styles
        for style in doc.styles:
            if style.type in [WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.CHARACTER]:
                self.extracted_styles[style.name] = self._extract_style_properties(style)
        
        # Analyze actual usage in document
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                style_name = paragraph.style.name
                if style_name not in self.extracted_styles:
                    self.extracted_styles[style_name] = {}
                
                # Extract paragraph-specific formatting
                self._extract_paragraph_formatting(paragraph, style_name)
        
        # Analyze tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            self._extract_paragraph_formatting(paragraph, 'table_content')
    
    def _extract_style_properties(self, style):
        """Extract comprehensive style properties"""
        properties = {
            'type': 'paragraph' if style.type == WD_STYLE_TYPE.PARAGRAPH else 'character',
            'font': {},
            'paragraph': {}
        }
        
        # Extract font properties
        if hasattr(style, 'font') and style.font:
            font = style.font
            if font.name:
                properties['font']['name'] = font.name
                self.font_families_used.add(font.name)
            if font.size:
                properties['font']['size'] = font.size.pt
            if font.bold is not None:
                properties['font']['bold'] = font.bold
            if font.italic is not None:
                properties['font']['italic'] = font.italic
            if font.underline is not None:
                properties['font']['underline'] = font.underline
            if font.color and font.color.rgb:
                color_hex = str(font.color.rgb)
                properties['font']['color'] = color_hex
                self.colors_used.add(color_hex)
        
        # Extract paragraph properties
        if style.type == WD_STYLE_TYPE.PARAGRAPH and hasattr(style, 'paragraph_format'):
            pf = style.paragraph_format
            if pf.alignment is not None:
                properties['paragraph']['alignment'] = pf.alignment
            if pf.space_before is not None:
                properties['paragraph']['space_before'] = pf.space_before.pt
            if pf.space_after is not None:
                properties['paragraph']['space_after'] = pf.space_after.pt
            if pf.line_spacing is not None:
                properties['paragraph']['line_spacing'] = pf.line_spacing
            if pf.first_line_indent is not None:
                properties['paragraph']['first_line_indent'] = pf.first_line_indent.pt
            if pf.left_indent is not None:
                properties['paragraph']['left_indent'] = pf.left_indent.pt
            if pf.right_indent is not None:
                properties['paragraph']['right_indent'] = pf.right_indent.pt
        
        return properties
    
    def _extract_paragraph_formatting(self, paragraph, style_name):
        """Extract formatting from actual paragraph usage"""
        if style_name not in self.extracted_styles:
            self.extracted_styles[style_name] = {'font': {}, 'paragraph': {}}
        
        # Extract paragraph format properties
        pf = paragraph.paragraph_format
        if pf.alignment is not None:
            self.extracted_styles[style_name]['paragraph']['alignment'] = pf.alignment
        if pf.space_before is not None:
            self.extracted_styles[style_name]['paragraph']['space_before'] = pf.space_before.pt
        if pf.space_after is not None:
            self.extracted_styles[style_name]['paragraph']['space_after'] = pf.space_after.pt
        if pf.left_indent is not None:
            self.extracted_styles[style_name]['paragraph']['left_indent'] = pf.left_indent.pt
        if pf.right_indent is not None:
            self.extracted_styles[style_name]['paragraph']['right_indent'] = pf.right_indent.pt
        if pf.first_line_indent is not None:
            self.extracted_styles[style_name]['paragraph']['first_line_indent'] = pf.first_line_indent.pt
        
        # Extract run-level formatting
        for run in paragraph.runs:
            if run.font.name:
                self.font_families_used.add(run.font.name)
            if run.font.color and run.font.color.rgb:
                self.colors_used.add(str(run.font.color.rgb))
    
    def _extract_document_properties(self, doc):
        """Extract document-level properties"""
        # Could extract document title, author, etc. from doc.core_properties
        pass
    
    def _initialize_comprehensive_latex(self):
        """Initialize LaTeX with comprehensive packages and settings"""
        self.latex_content = [
            f'\\documentclass[11pt,a4paper]{{{self.document_class}}}',
            '',
            '% Essential packages',
            '\\usepackage[utf8]{inputenc}',
            '\\usepackage[T1]{fontenc}',
            '\\usepackage[english]{babel}',
            '\\usepackage{geometry}',
            '',
            '% Math packages',
            '\\usepackage{amsmath}',
            '\\usepackage{amsfonts}',
            '\\usepackage{amssymb}',
            '\\usepackage{mathtools}',
            '',
            '% Graphics and colors',
            '\\usepackage{graphicx}',
            '\\usepackage{xcolor}',
            '\\usepackage{colortbl}',
            '',
            '% Table packages',
            '\\usepackage{booktabs}',
            '\\usepackage{array}',
            '\\usepackage{longtable}',
            '\\usepackage{tabularx}',
            '\\usepackage{multirow}',
            '',
            '% List packages',
            '\\usepackage{enumerate}',
            '\\usepackage{enumitem}',
            '',
            '% Text formatting',
            '\\usepackage{soul}',
            '\\usepackage{ulem}',
            '\\usepackage{textcomp}',
            '',
            '% Layout and spacing',
            '\\usepackage{setspace}',
            '\\usepackage{indentfirst}',
            '\\usepackage{changepage}',
            '',
            '% Hyperlinks',
            '\\usepackage{hyperref}',
            '',
            '% Font packages',
        ]
        
        # Add font support
        self._add_font_support()
        
        # Add color definitions
        self._add_color_definitions()
        
        # Add geometry settings
        self.latex_content.extend([
            '',
            '% Document geometry',
            '\\geometry{margin=1in}',
            '',
            '% Paragraph settings',
            '\\setlength{\\parindent}{0pt}',
            '\\setlength{\\parskip}{6pt plus 2pt minus 1pt}',
            '',
            '% Custom commands',
        ])
        
        # Add custom commands
        self._add_custom_commands()
        
        self.latex_content.extend([
            '',
            '\\begin{document}',
            ''
        ])
    
    def _add_font_support(self):
        """Add font support based on fonts used in document"""
        if 'Times New Roman' in self.font_families_used:
            self.latex_content.append('\\usepackage{times}')
        if 'Arial' in self.font_families_used or 'Helvetica' in self.font_families_used:
            self.latex_content.append('\\usepackage{helvet}')
        if 'Courier' in self.font_families_used or 'Courier New' in self.font_families_used:
            self.latex_content.append('\\usepackage{courier}')
    
    def _add_color_definitions(self):
        """Add color definitions based on colors used"""
        if self.colors_used:
            self.latex_content.append('')
            self.latex_content.append('% Color definitions')
            for i, color_hex in enumerate(sorted(self.colors_used)):
                if color_hex != 'None':
                    try:
                        # Convert hex to RGB
                        color_hex = color_hex.replace('#', '')
                        if len(color_hex) == 6:
                            r = int(color_hex[0:2], 16) / 255.0
                            g = int(color_hex[2:4], 16) / 255.0
                            b = int(color_hex[4:6], 16) / 255.0
                            self.latex_content.append(f'\\definecolor{{customcolor{i}}}{{rgb}}{{{r:.3f},{g:.3f},{b:.3f}}}')
                    except:
                        pass
    
    def _add_custom_commands(self):
        """Add custom LaTeX commands for consistent formatting"""
        self.latex_content.extend([
            '% Custom formatting commands',
            '\\newcommand{\\customspacing}[1]{\\vspace{#1}}',
            '\\newcommand{\\customindent}[1]{\\hspace{#1}}',
        ])
    
    def _process_document_content(self, doc):
        """Process document content with comprehensive formatting"""
        # Process main paragraphs
        self._process_paragraphs_advanced(doc.paragraphs)
        
        # Process tables with advanced formatting
        self._process_tables_advanced(doc.tables)
    
    def _process_paragraphs_advanced(self, paragraphs):
        """Process paragraphs with advanced formatting preservation"""
        for paragraph in paragraphs:
            if paragraph.text.strip():
                latex_para = self._convert_paragraph_advanced(paragraph)
                if latex_para:
                    self.latex_content.extend(latex_para)
                    self.latex_content.append('')
    
    def _convert_paragraph_advanced(self, paragraph):
        """Convert paragraph with comprehensive formatting"""
        text = paragraph.text.strip()
        if not text:
            return []
        
        style_name = paragraph.style.name
        result = []
        
        # Handle special styles first
        if style_name in self.style_mappings:
            content = self._process_runs_advanced(paragraph.runs)
            formatted_content = self.style_mappings[style_name].format(content=content)
            return [formatted_content]
        
        # Handle lists with proper formatting
        if self._is_advanced_list_item(paragraph):
            return self._convert_advanced_list_item(paragraph)
        
        # Process regular paragraph with formatting
        paragraph_latex = []
        
        # Add paragraph formatting
        para_formatting = self._get_paragraph_formatting(paragraph)
        if para_formatting['start']:
            paragraph_latex.extend(para_formatting['start'])
        
        # Process the paragraph content
        content = self._process_runs_advanced(paragraph.runs)
        if content.strip():
            paragraph_latex.append(content)
        
        # Add paragraph formatting end
        if para_formatting['end']:
            paragraph_latex.extend(para_formatting['end'])
        
        return paragraph_latex
    
    def _get_paragraph_formatting(self, paragraph):
        """Extract comprehensive paragraph formatting"""
        formatting = {'start': [], 'end': []}
        pf = paragraph.paragraph_format
        
        # Handle alignment
        if pf.alignment and pf.alignment in self.alignment_map:
            env = self.alignment_map[pf.alignment]
            formatting['start'].append(f'\\begin{{{env}}}')
            formatting['end'].insert(0, f'\\end{{{env}}}')
        
        # Handle indentation
        indent_commands = []
        if pf.left_indent and pf.left_indent.pt > 0:
            indent_commands.append(f'\\customindent{{{pf.left_indent.pt}pt}}')
        if pf.first_line_indent and pf.first_line_indent.pt != 0:
            if pf.first_line_indent.pt > 0:
                indent_commands.append(f'\\indent\\customindent{{{pf.first_line_indent.pt}pt}}')
            else:
                indent_commands.append('\\noindent')
        
        if indent_commands:
            formatting['start'].extend(indent_commands)
        
        # Handle spacing
        if pf.space_before and pf.space_before.pt > 0:
            formatting['start'].append(f'\\customspacing{{{pf.space_before.pt}pt}}')
        if pf.space_after and pf.space_after.pt > 0:
            formatting['end'].append(f'\\customspacing{{{pf.space_after.pt}pt}}')
        
        return formatting
    
    def _process_runs_advanced(self, runs):
        """Process runs with comprehensive formatting"""
        content_parts = []
        
        for run in runs:
            text = run.text
            if not text:
                continue
            
            # Escape LaTeX special characters
            escaped_text = self._escape_latex_chars(text)
            
            # Apply comprehensive formatting
            formatted_text = self._apply_comprehensive_run_formatting(escaped_text, run)
            content_parts.append(formatted_text)
        
        return ''.join(content_parts)
    
    def _apply_comprehensive_run_formatting(self, text, run):
        """Apply comprehensive formatting to text"""
        if not text.strip():
            return text
        
        formatted = text
        
        # Font family
        if run.font.name and run.font.name != 'Calibri':  # Don't change default
            if run.font.name == 'Times New Roman':
                formatted = f'\\textrm{{{formatted}}}'
            elif run.font.name in ['Arial', 'Helvetica']:
                formatted = f'\\textsf{{{formatted}}}'
            elif run.font.name in ['Courier', 'Courier New']:
                formatted = f'\\texttt{{{formatted}}}'
        
        # Font size
        if run.font.size:
            size_pt = run.font.size.pt
            if size_pt >= 24:
                formatted = f'\\Huge{{{formatted}}}'
            elif size_pt >= 20:
                formatted = f'\\huge{{{formatted}}}'
            elif size_pt >= 17:
                formatted = f'\\LARGE{{{formatted}}}'
            elif size_pt >= 14:
                formatted = f'\\Large{{{formatted}}}'
            elif size_pt >= 12:
                formatted = f'\\large{{{formatted}}}'
            elif size_pt >= 10:
                formatted = f'\\normalsize{{{formatted}}}'
            elif size_pt >= 9:
                formatted = f'\\small{{{formatted}}}'
            elif size_pt >= 8:
                formatted = f'\\footnotesize{{{formatted}}}'
            elif size_pt >= 6:
                formatted = f'\\scriptsize{{{formatted}}}'
            else:
                formatted = f'\\tiny{{{formatted}}}'
        
        # Color
        if run.font.color and run.font.color.rgb:
            color_hex = str(run.font.color.rgb)
            if color_hex in self.colors_used:
                color_index = sorted(list(self.colors_used)).index(color_hex)
                formatted = f'\\textcolor{{customcolor{color_index}}}{{{formatted}}}'
        
        # Bold, italic, underline (apply in reverse order to nest properly)
        if run.underline:
            formatted = f'\\underline{{{formatted}}}'
        if run.italic:
            formatted = f'\\textit{{{formatted}}}'
        if run.bold:
            formatted = f'\\textbf{{{formatted}}}'
        
        return formatted
    
    def _is_advanced_list_item(self, paragraph):
        """Advanced list item detection"""
        text = paragraph.text.strip()
        # Check for various list patterns
        return bool(re.match(r'^[•\-\*]\s+|^\d+[\.)]\s+|^[a-zA-Z][\.)]\s+|^[ivxlcdm]+[\.)]\s+', text, re.IGNORECASE))
    
    def _convert_advanced_list_item(self, paragraph):
        """Convert list item with proper formatting"""
        text = paragraph.text.strip()
        
        # Determine list type and extract content
        if re.match(r'^\d+[\.)]\s+', text):
            # Numbered list
            content = re.sub(r'^\d+[\.)]\s+', '', text)
            escaped_content = self._process_runs_advanced(paragraph.runs)
            return [
                '\\begin{enumerate}',
                f'\\item {escaped_content}',
                '\\end{enumerate}'
            ]
        elif re.match(r'^[a-zA-Z][\.)]\s+', text):
            # Alphabetic list
            content = re.sub(r'^[a-zA-Z][\.)]\s+', '', text)
            escaped_content = self._process_runs_advanced(paragraph.runs)
            return [
                '\\begin{enumerate}[label=\\alph*.]',
                f'\\item {escaped_content}',
                '\\end{enumerate}'
            ]
        elif re.match(r'^[ivxlcdm]+[\.)]\s+', text, re.IGNORECASE):
            # Roman numeral list
            content = re.sub(r'^[ivxlcdm]+[\.)]\s+', '', text, flags=re.IGNORECASE)
            escaped_content = self._process_runs_advanced(paragraph.runs)
            return [
                '\\begin{enumerate}[label=\\roman*.]',
                f'\\item {escaped_content}',
                '\\end{enumerate}'
            ]
        else:
            # Bulleted list
            content = re.sub(r'^[•\-\*]\s+', '', text)
            escaped_content = self._process_runs_advanced(paragraph.runs)
            return [
                '\\begin{itemize}',
                f'\\item {escaped_content}',
                '\\end{itemize}'
            ]
    
    def _process_tables_advanced(self, tables):
        """Process tables with advanced formatting"""
        for table in tables:
            latex_table = self._convert_table_advanced(table)
            if latex_table:
                self.latex_content.extend(latex_table)
                self.latex_content.append('')
    
    def _convert_table_advanced(self, table):
        """Convert table with comprehensive formatting"""
        if not table.rows:
            return []
        
        # Determine number of columns
        max_cols = max(len(row.cells) for row in table.rows)
        
        # Create column specification with proper alignment
        col_spec = 'l' * max_cols  # Default to left alignment
        
        # Start table with advanced formatting
        latex_lines = [
            '\\begin{table}[h]',
            '\\centering',
            f'\\begin{{tabular}}{{{col_spec}}}',
            '\\toprule'
        ]
        
        # Process rows with formatting
        for i, row in enumerate(table.rows):
            row_content = []
            for cell in row.cells[:max_cols]:
                # Process cell content with formatting
                cell_latex = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        para_content = self._process_runs_advanced(paragraph.runs)
                        if para_content.strip():
                            cell_latex.append(para_content)
                
                cell_text = ' '.join(cell_latex) if cell_latex else ''
                row_content.append(cell_text)
            
            # Pad row if necessary
            while len(row_content) < max_cols:
                row_content.append('')
            
            latex_lines.append(' & '.join(row_content) + ' \\\\')
            
            # Add midrule after first row (header)
            if i == 0 and len(table.rows) > 1:
                latex_lines.append('\\midrule')
        
        # End table
        latex_lines.extend([
            '\\bottomrule',
            '\\end{tabular}',
            '\\end{table}'
        ])
        
        return latex_lines
    
    def _escape_latex_chars(self, text):
        """Comprehensive LaTeX character escaping"""
        # Enhanced escape mappings
        escape_chars = {
            '\\': r'\textbackslash{}',
            '{': r'\{',
            '}': r'\}',
            '$': r'\$',
            '&': r'\&',
            '%': r'\%',
            '#': r'\#',
            '^': r'\textasciicircum{}',
            '_': r'\_',
            '~': r'\textasciitilde{}',
            '"': r'\'\'',
            '`': r'\`{}',
            ''': r'\textquoteleft{}',
            ''': r'\textquoteright{}',
            '"': r'\textquotedblleft{}',
            '"': r'\textquotedblright{}',
            '–': r'--',
            '—': r'---',
            '…': r'\ldots{}',
        }
        
        for char, escaped in escape_chars.items():
            text = text.replace(char, escaped)
        
        return text
    
    def _finalize_latex_document(self):
        """Finalize LaTeX document with proper ending"""
        self.latex_content.extend([
            '',
            '\\end{document}'
        ])

class DocumentFormatter:
    # Pre-compiled regex patterns for better performance
    HEADING_PATTERN = re.compile(r'^\d+\.|^[A-Z][^.!?]*$')
    LIST_PATTERN = re.compile(r'^[•\-\*]|^\d+[.).]')
    
    # Content type constants
    CONTENT_TYPES = {
        'heading': 'heading',
        'list_item': 'list_item', 
        'title': 'title',
        'quote': 'quote',
        'body': 'body'
    }
    
    # Built-in styles that shouldn't be recreated
    BUILTIN_STYLES = frozenset([
        'Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 
        'Heading 5', 'Heading 6', 'Title', 'Subtitle', 'Header', 'Footer'
    ])

    def __init__(self, debug=False):
        self.template_styles = {}
        self.style_content_map = {}
        self.content_style_cache = {}
        self.debug = debug
        if debug:
            logger.setLevel(logging.INFO)
    
    def extract_styles_from_template(self, template_path):
        """Extract styles and analyze content patterns efficiently"""
        doc = Document(template_path)
        
        # Single pass through styles for extraction
        paragraph_styles = {}
        character_styles = {}
        
        for style in doc.styles:
            style_data = self._extract_style_data(style)
            if style_data:
                if style.type == WD_STYLE_TYPE.PARAGRAPH:
                    paragraph_styles[style.name] = style_data
                elif style.type == WD_STYLE_TYPE.CHARACTER:
                    character_styles[style.name] = style_data
        
        self.template_styles = {**paragraph_styles, **character_styles}
        
        # Single pass content analysis with batch processing
        self._analyze_content_usage_batch(doc)
        
        if self.debug:
            logger.info(f"Extracted {len(self.template_styles)} styles")
    
    def _extract_style_data(self, style):
        """Efficiently extract style data with minimal attribute access"""
        style_data = {
            'type': 'paragraph' if style.type == WD_STYLE_TYPE.PARAGRAPH else 'character',
            'name': style.name
        }
        
        # Extract font data efficiently
        font_data = {}
        if hasattr(style, 'font') and style.font:
            font = style.font
            font_attrs = ['name', 'size', 'bold', 'italic', 'underline']
            for attr in font_attrs:
                value = getattr(font, attr, None)
                if value is not None:
                    font_data[attr] = value
            
            # Handle color separately due to complex access pattern
            if hasattr(font, 'color') and font.color and hasattr(font.color, 'rgb') and font.color.rgb:
                font_data['color'] = font.color.rgb
        
        if font_data:
            style_data['font'] = font_data
        
        # Extract paragraph data efficiently for paragraph styles
        if style.type == WD_STYLE_TYPE.PARAGRAPH and hasattr(style, 'paragraph_format') and style.paragraph_format:
            para_data = {}
            pf = style.paragraph_format
            para_attrs = ['alignment', 'space_before', 'space_after', 'line_spacing', 
                         'first_line_indent', 'left_indent', 'right_indent']
            
            for attr in para_attrs:
                value = getattr(pf, attr, None)
                if value is not None:
                    para_data[attr] = value
            
            if para_data:
                style_data['paragraph'] = para_data
        
        return style_data if len(style_data) > 2 else None  # Only return if has meaningful data
    
    def _analyze_content_usage_batch(self, doc):
        """Batch analyze content patterns for optimal performance"""
        style_usage = defaultdict(list)
        
        # Collect all text content in single pass
        all_paragraphs = []
        all_paragraphs.extend(doc.paragraphs)
        
        # Add table paragraphs
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
        
        # Process all paragraphs with batch categorization
        content_types = []
        texts = []
        style_names = []
        
        for paragraph in all_paragraphs:
            text = paragraph.text.strip()
            if text and len(text) > 2:  # Skip very short text
                texts.append(text[:100])  # Limit text length for performance
                style_names.append(paragraph.style.name)
                content_types.append(self._categorize_content_fast(text))
        
        # Build style mappings efficiently
        for style_name, content_type in zip(style_names, content_types):
            if style_name in self.template_styles:
                if style_name not in self.style_content_map:
                    self.style_content_map[style_name] = []
                self.style_content_map[style_name].append(content_type)
        
        # Determine primary content type for each style
        for style_name, content_list in self.style_content_map.items():
            if content_list:
                primary_type = Counter(content_list).most_common(1)[0][0]
                self.template_styles[style_name]['primary_content_type'] = primary_type
    
    @lru_cache(maxsize=1000)
    def _categorize_content_fast(self, text):
        """Fast content categorization with caching and optimized patterns"""
        text_stripped = text.strip()
        text_len = len(text_stripped)
        
        # Quick length-based filtering
        if text_len == 0:
            return self.CONTENT_TYPES['body']
        
        # Check for titles (short uppercase)
        if text_stripped.isupper() and text_len < 50:
            return self.CONTENT_TYPES['title']
        
        # Check for quotes (starts with quote marks)
        first_char = text_stripped[0]
        if first_char in '\"\'':
            return self.CONTENT_TYPES['quote']
        
        # Use pre-compiled patterns for better performance
        if text_len < 100 and self.HEADING_PATTERN.match(text_stripped):
            return self.CONTENT_TYPES['heading']
        
        if self.LIST_PATTERN.match(text_stripped):
            return self.CONTENT_TYPES['list_item']
        
        return self.CONTENT_TYPES['body']
    
    def apply_styles_to_target(self, target_path):
        """Apply styles with optimized batch processing"""
        target_doc = Document(target_path)
        
        # Pre-create style mapping for fast lookup
        self._prepare_style_mapping()
        
        # Batch update existing styles
        self._batch_update_styles(target_doc)
        
        # Batch apply styles to content
        self._batch_apply_paragraph_styles(target_doc)
        
        return target_doc
    
    def _prepare_style_mapping(self):
        """Pre-compute style mappings for O(1) lookup"""
        content_type_styles = defaultdict(list)
        
        for style_name, style_data in self.template_styles.items():
            content_type = style_data.get('primary_content_type', 'body')
            content_type_styles[content_type].append(style_name)
        
        # Cache the best style for each content type
        for content_type, style_list in content_type_styles.items():
            if style_list:
                self.content_style_cache[content_type] = style_list[0]
        
        # Add fallback mappings
        style_name_lower_map = {name.lower(): name for name in self.template_styles.keys()}
        
        for content_type in self.CONTENT_TYPES.values():
            if content_type not in self.content_style_cache:
                # Try to find by name pattern
                for pattern in [content_type, 'heading', 'title', 'normal']:
                    if pattern in style_name_lower_map:
                        self.content_style_cache[content_type] = style_name_lower_map[pattern]
                        break
        
        # Ultimate fallback
        if self.template_styles:
            fallback_style = next(iter(self.template_styles.keys()))
            for content_type in self.CONTENT_TYPES.values():
                if content_type not in self.content_style_cache:
                    self.content_style_cache[content_type] = fallback_style
    
    def _batch_update_styles(self, target_doc):
        """Efficiently update or create styles in batch"""
        existing_styles = {style.name for style in target_doc.styles}
        
        for style_name, style_data in self.template_styles.items():
            try:
                if style_name in existing_styles:
                    self._update_style_fast(target_doc.styles[style_name], style_data)
                elif style_name not in self.BUILTIN_STYLES:
                    self._create_style_fast(target_doc, style_name, style_data)
            except Exception as e:
                if self.debug:
                    logger.warning(f"Style operation failed for {style_name}: {e}")
    
    def _update_style_fast(self, style, style_data):
        """Fast style update with minimal attribute access"""
        # Update font properties
        if 'font' in style_data and hasattr(style, 'font') and style.font:
            font_data = style_data['font']
            font = style.font
            
            # Batch update font attributes
            for attr, value in font_data.items():
                if attr == 'color':
                    if hasattr(font, 'color') and font.color:
                        font.color.rgb = value
                else:
                    setattr(font, attr, value)
        
        # Update paragraph properties
        if 'paragraph' in style_data and hasattr(style, 'paragraph_format') and style.paragraph_format:
            para_data = style_data['paragraph']
            pf = style.paragraph_format
            
            # Batch update paragraph attributes
            for attr, value in para_data.items():
                setattr(pf, attr, value)
    
    def _create_style_fast(self, doc, style_name, style_data):
        """Fast style creation with error handling"""
        try:
            style_type = WD_STYLE_TYPE.PARAGRAPH if style_data['type'] == 'paragraph' else WD_STYLE_TYPE.CHARACTER
            new_style = doc.styles.add_style(style_name, style_type)
            self._update_style_fast(new_style, style_data)
        except Exception as e:
            if self.debug:
                logger.warning(f"Failed to create style {style_name}: {e}")
    
    def _batch_apply_paragraph_styles(self, target_doc):
        """Apply styles to all paragraphs in optimized batch"""
        # Collect all paragraphs
        all_paragraphs = list(target_doc.paragraphs)
        
        # Add table paragraphs
        for table in target_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)
        
        # Batch process paragraphs
        for paragraph in all_paragraphs:
            text = paragraph.text.strip()
            if text:
                content_type = self._categorize_content_fast(text)
                best_style = self.content_style_cache.get(content_type)
                
                if best_style:
                    try:
                        paragraph.style = target_doc.styles[best_style]
                        if self.debug:
                            logger.info(f"Applied '{best_style}' to: {text[:30]}...")
                    except KeyError:
                        if self.debug:
                            logger.warning(f"Style '{best_style}' not found")
    
    def apply_formatting(self, template_path, target_path):
        """Optimized main formatting method"""
        if self.debug:
            logger.info(f"Processing: {template_path} -> {target_path}")
        
        # Extract styles and patterns
        self.extract_styles_from_template(template_path)
        
        # Apply to target document
        target_doc = self.apply_styles_to_target(target_path)
        
        # Save efficiently
        output_path = tempfile.mktemp(suffix='.docx')
        target_doc.save(output_path)
        
        if self.debug:
            logger.info(f"Completed: {output_path}")
        
        return output_path
    
    def convert_to_latex(self, docx_path):
        """Convert Word document to LaTeX with comprehensive formatting"""
        converter = LaTeXConverter(debug=self.debug)
        return converter.convert_document(docx_path)
    
    def __del__(self):
        """Clean up cache when object is destroyed"""
        self._categorize_content_fast.cache_clear()